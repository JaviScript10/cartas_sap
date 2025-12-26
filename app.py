import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from io import StringIO

# ================= CONFIG =================
st.set_page_config(
    page_title="Automatizador de Cartas",
    page_icon="⚡",
    layout="wide"
)

# 1. PRIMERO INICIALIZAMOS (Esto debe ir antes de cualquier 'if')
if 'count_reset' not in st.session_state:
    st.session_state.count_reset = 0

if 'limpieza_inicial' not in st.session_state:
    st.session_state.limpieza_inicial = False

# 2. SEGUIDO HACEMOS LA VALIDACIÓN DE LIMPIEZA
if st.session_state.count_reset == 0 and not st.session_state.limpieza_inicial:
    st.session_state.limpieza_inicial = True
    st.rerun()

if 'carta_generada' not in st.session_state:
    st.session_state.carta_generada = False

if 'output_path' not in st.session_state:
    st.session_state.output_path = None

if 'vista_previa_html' not in st.session_state:
    st.session_state.vista_previa_html = None

os.makedirs("templates", exist_ok=True)
os.makedirs("output", exist_ok=True)

# ================= GERENTES POR ZONA =================
GERENTES = {
    "Norte": "Christian Alberto Gómez Díaz",
    "Centro": "Alex Andrés González Villablanca",
    "Sur": "Christian Enrique Araya Silva"
}

# ================= CATEGORÍAS Y TIPOS DE CARTAS =================
CATEGORIAS = {
    "Cobros": {
        "apertura_casa_halu": "Apertura Casa Cerrada NOLU",
        "carta_aporte_lectura": "Aporte Lectura",
        "atencion_emergencia_halu": "Atención de Emergencias HALU",
        "aumento_consumo_halu_sinvisita": "Aumento Consumo HALU Sin visita técnica",
        "aumento_consumo_nolu_sinvisita": "Aumento Consumo NOLU Sin visita técnica",
        "carta_compromiso": "Carta Compromiso 10 días",
        "carta_compromiso_i5": "Carta Compromiso con I5",
        "carta_falta_info": "Carta falta información reclamo",
        "error_lectura_halu": "Error de Lectura HALU",
        "error_lectura_nolu": "Error de Lectura NOLU",
        "error_lectura_regularizado_sgte_lectura": "Error Lectura Regularizado Siguiente Lectura HALU",
        "facturaciones_normalizadas": "Facturaciones Normalizadas",
        "normal_avance": "Normal Avance"
    },
    "DAR (Artefacto Dañado)": {
        # Por ahora vacío, se agregarán después
    },
    "Técnico Comercial": {
        # Por ahora vacío, se agregarán después
    }
}

# ================= CANALES CON GÉNERO =================
CANALES_MASCULINO = ["Call Center", "Correo Electrónico"]
CANALES_FEMENINO = ["Oficina Comercial", "Página Web", "App CGE 1Click"]
CANALES_EXTERNOS = ["Portal SEC", "Portal SERNAC"]  # Sin "nuestro/nuestra"

# ================= FUNCIONES AUXILIARES =================
def formatear_monto(valor):
    """Formatea monto con $ y puntos de miles"""
    try:
        valor_limpio = str(valor).replace('$', '').replace('.', '').replace(',', '').strip()
        if valor_limpio:
            numero = int(valor_limpio)
            return f"${numero:,}".replace(',', '.')
        return valor
    except:
        return valor

def capitalizar_texto(texto):
    """
    Capitaliza la primera letra de cada palabra
    eduardo lópez → Eduardo López
    """
    if texto:
        return texto.title()
    return texto

def formatear_fecha(fecha_input):
    """
    Formatea fecha de DDMMYYYY a DD/MM/YYYY
    20022025 → 20/02/2025
    También acepta ya formateadas: 20/02/2025 → 20/02/2025
    """
    if not fecha_input:
        return fecha_input
    
    # Limpiar entrada (quitar espacios, guiones, puntos, barras)
    fecha_limpia = str(fecha_input).replace('/', '').replace('-', '').replace('.', '').replace(' ', '').strip()
    
    # Si tiene exactamente 8 dígitos, formatear
    if len(fecha_limpia) == 8 and fecha_limpia.isdigit():
        dia = fecha_limpia[0:2]
        mes = fecha_limpia[2:4]
        anio = fecha_limpia[4:8]
        
        # Validación básica
        try:
            dia_int = int(dia)
            mes_int = int(mes)
            anio_int = int(anio)
            
            if 1 <= dia_int <= 31 and 1 <= mes_int <= 12 and 2000 <= anio_int <= 2100:
                return f"{dia}/{mes}/{anio}"
        except:
            pass
    
    # Si ya está formateada o no es válida, devolver tal cual
    return fecha_input

# ================= HEADER =================
st.markdown("""
<div style='background:linear-gradient(90deg,#1e40af,#0ea5e9);
padding:25px;border-radius:12px;margin-bottom:25px;box-shadow:0 4px 6px rgba(0,0,0,0.1)'>
<h1 style='color:white;margin:0;font-size:2.2em'>⚡ AUTOMATIZADOR DE CARTAS </h1>
<p style='color:#e0f2fe;margin:8px 0 0 0;font-size:1.1em'>Sistema de Generación Automática de Respuestas</p>
</div>
""", unsafe_allow_html=True)

# ================= SELECTOR DE CATEGORÍA Y CARTA =================
st.markdown("### 📋 SELECCIONAR TIPO DE CARTA")

# Crear dos columnas: una para los selectores (columna), otra para el botón
col_selectores, col_boton = st.columns([2, 1])

with col_selectores:
    # Selector de Categoría
    categoria = st.selectbox(
        "Categoría:",
        options=list(CATEGORIAS.keys()),
        key=f"categoria_{st.session_state.count_reset}"
    )
    
    # Mostrar info si la categoría está vacía
    if not CATEGORIAS[categoria]:
        st.info(f"ℹ️ La categoría '{categoria}' aún no tiene cartas disponibles. Se agregarán próximamente.")
        st.stop()
    
    # Selector de Carta
    tipo_carta = st.selectbox(
        "Tipo de carta:",
        options=list(CATEGORIAS[categoria].keys()),
        format_func=lambda x: CATEGORIAS[categoria][x],
        key=f"tipo_carta_{st.session_state.count_reset}"
    )

with col_boton:
    # Espaciador para alinear verticalmente con la categoría
    st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)
    
    # Estilo para el botón rectangular - BLANCO con NEGRO
    st.markdown("""
        <style>
        div[data-testid="stButton"] button[kind="secondary"] {
            background-color: white !important;
            background: white !important;
            color: black !important;
            border: 2px solid #000000 !important;
            font-weight: 700 !important;
            font-size: 0.9rem !important;
            border-radius: 8px !important;
            padding: 12px 8px !important;
            box-shadow: 0 2px 6px rgba(0, 0, 0, 0.2) !important;
            transition: all 0.3s !important;
            width: 100% !important;
            height: 60px !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            text-align: center !important;
            white-space: normal !important;
            line-height: 1.2 !important;
        }
        div[data-testid="stButton"] button[kind="secondary"]:hover {
            background-color: #f0f0f0 !important;
            background: #f0f0f0 !important;
            box-shadow: 0 4px 10px rgba(0, 0, 0, 0.3) !important;
            transform: scale(1.02) !important;
            border-color: #000000 !important;
            color: black !important;
        }
        div[data-testid="stButton"] button[kind="secondary"]:active {
            transform: scale(0.98) !important;
            background-color: #e0e0e0 !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Botón rectangular con texto
    if st.button("🔄 REINICIAR FORMULARIO", type="secondary", use_container_width=True, key=f"btn_reset_{st.session_state.count_reset}"):
        st.session_state.count_reset += 1 
        st.session_state.carta_generada = False
        st.session_state.output_path = None
        st.session_state.vista_previa_html = None
        st.rerun()

st.markdown("---")

# ================= FORMULARIO =================
col1, col2 = st.columns(2)

with col1:
    st.markdown("### 📋 DATOS DEL CLIENTE")
    
    # ⭐ MEJORA: Auto-capitalizar comuna
    comuna_raw = st.text_input(
        "Comuna:",
        value="",
        placeholder="Ingrese comuna (Ej: Valparaíso)", 
        key=f"comuna_input_{st.session_state.count_reset}",
        autocomplete="new-password"
    )
    comuna = capitalizar_texto(comuna_raw)
    
    # ⭐ NUEVO NOMBRE: Formalidad en lugar de Tratamiento
    tratamiento = st.selectbox(
        "Formalidad (Señor o Señora):", 
        ["", "Señor", "Señora"],
        key=f"tratamiento_{st.session_state.count_reset}",
        help="Forma de dirigirse al cliente"
    )
    
    # ⭐ MEJORA: Auto-capitalizar nombre
    nombre_cliente_raw = st.text_input(
        "Nombre y apellido completo:",
        placeholder="Eduardo López",
        key=f"nombre_{st.session_state.count_reset}"
    )
    nombre_cliente = capitalizar_texto(nombre_cliente_raw)
    
    # ⭐ MEJORA: Auto-capitalizar dirección
    direccion_raw = st.text_input(
        "Dirección:",
        placeholder="Prat 725",
        key=f"direccion_{st.session_state.count_reset}"
    )
    direccion = capitalizar_texto(direccion_raw)
    
    numero_cliente = st.text_input(
        "Número cliente:",
        placeholder="6255126",
        key=f"num_cliente_{st.session_state.count_reset}",
        autocomplete="off"
    )

with col2:
    st.markdown("### 📄 DATOS DEL RECLAMO")
    
    gr_numero = st.text_input(
        "N° GR (Número de Reclamo):",
        placeholder="15624563",
        key=f"gr_{st.session_state.count_reset}",
        help="Este número se usará en DGR y en Ref.: Reclamo N°",
        autocomplete="off"
    )
    
    # ⭐ NUEVO: Campo opcional para casos SEC/SERNAC
    tiene_caso_sec = st.checkbox(
        "¿Es un caso SEC/SERNAC?",
        key=f"tiene_caso_sec_{st.session_state.count_reset}",
        help="Marcar si el reclamo viene de Portal SEC o Portal SERNAC"
    )
    
    caso_sec_numero = ""
    tipo_caso = ""
    if tiene_caso_sec:
        col_sec1, col_sec2 = st.columns(2)
        
        with col_sec1:
            tipo_caso = st.selectbox(
                "Tipo:",
                ["SEC", "SERNAC"],
                key=f"tipo_caso_{st.session_state.count_reset}",
                help="Seleccione si es SEC o SERNAC"
            )
        
        with col_sec2:
            caso_sec_numero = st.text_input(
                f"N° Caso {tipo_caso}:",
                placeholder="123456",
                key=f"caso_sec_numero_{st.session_state.count_reset}",
                help=f"Número del caso en {tipo_caso}"
            )
    
    zona = st.selectbox(
        "Firma - Zona Geográfica:",
        ["Norte", "Centro", "Sur"],
        help="Seleccione la zona para el gerente comercial correspondiente",
        key=f"zona_{st.session_state.count_reset}"
    )
    
    st.info(f"👤 Gerente Comercial: {GERENTES[zona]}")
    
    canal_ingreso = st.selectbox(
        "Canal de ingreso del reclamo:",
        [
            "Oficina Comercial",
            "WhatsApp",
            "App CGE 1Click",
            "Call Center",
            "Correo Electrónico",
            "Página Web",
            "Portal SEC",
            "Portal SERNAC"
        ],
        key=f"canal_{st.session_state.count_reset}"
    )

st.markdown("---")

# ================= DATOS ESPECÍFICOS POR TIPO DE CARTA =================
if tipo_carta == "error_lectura_halu":
    with st.expander("📊 DATOS ADICIONALES DE FACTURACIÓN (Opcional)"):
        col_a, col_b, col_c = st.columns(3)
        
        with col_a:
            fecha_boleta_raw = st.text_input(
                "Fecha boleta:", 
                placeholder="20122025 o 20/12/2025",
                key=f"fecha_boleta_{st.session_state.count_reset}",
                help="Formato: DDMMYYYY (ej: 20122025) o DD/MM/YYYY"
            )
            fecha_boleta = formatear_fecha(fecha_boleta_raw)
            if fecha_boleta_raw and fecha_boleta != fecha_boleta_raw:
                st.caption(f"📅 Formateado: {fecha_boleta}")
            tipo_doc = st.selectbox(
                "Tipo documento:", 
                ["boleta", "factura"],
                key=f"tipo_doc_{st.session_state.count_reset}"
            )
            numero_boleta = st.text_input(
                "N° Boleta/Factura:", 
                placeholder="123456",
                key=f"num_boleta_{st.session_state.count_reset}",
                help="Número de la boleta o factura generada"
            )
        
        with col_b:
            consumo_kwh = st.text_input(
                "Consumo kWh:", 
                placeholder="350",
                key=f"consumo_kwh_{st.session_state.count_reset}"
            )
            monto_boleta_input = st.text_input(
                "Monto:", 
                placeholder="45000",
                key=f"monto_{st.session_state.count_reset}"
            )
            monto_boleta = formatear_monto(monto_boleta_input) if monto_boleta_input else ""
            if monto_boleta:
                st.caption(f"💰 Formateado: {monto_boleta}")
        
        with col_c:
            dia_inicio = st.text_input(
                "Día inicio lectura:", 
                placeholder="06",
                key=f"dia_inicio_{st.session_state.count_reset}"
            )
            dia_fin = st.text_input(
                "Día fin lectura:", 
                placeholder="12",
                key=f"dia_fin_{st.session_state.count_reset}"
            )
            rango_lectura = f"{dia_inicio} y {dia_fin}" if dia_inicio and dia_fin else ""

elif tipo_carta == "apertura_casa_halu":
    with st.expander("📊 DATOS ESPECÍFICOS - APERTURA CASA CERRADA", expanded=True):
        st.markdown("#### 📅 Periodo sin acceso al medidor")
        
        col_a, col_b, col_c, col_d = st.columns(4)
        
        with col_a:
            periodo_inicio = st.text_input(
                "Mes inicio sin acceso:",
                placeholder="marzo",
                key=f"periodo_inicio_{st.session_state.count_reset}",
                help="Ejemplo: marzo, abril, mayo..."
            )
        
        with col_b:
            anio_inicio = st.text_input(
                "Año inicio:",
                placeholder="2024",
                value="2024",
                key=f"anio_inicio_{st.session_state.count_reset}",
                help="Año en que comenzó el periodo sin acceso"
            )
        
        with col_c:
            periodo_fin = st.text_input(
                "Mes fin sin acceso:",
                placeholder="agosto",
                key=f"periodo_fin_{st.session_state.count_reset}",
                help="Ejemplo: agosto, septiembre..."
            )
        
        with col_d:
            anio_fin = st.text_input(
                "Año fin:",
                placeholder="2025",
                value="2025",
                key=f"anio_fin_{st.session_state.count_reset}",
                help="Año en que se accedió al medidor"
            )
        
        st.markdown("---")
        st.markdown("#### 🔢 Número de medidor")
        
        numero_medidor = st.text_input(
            "N° Medidor:",
            placeholder="E044124",
            key=f"numero_medidor_{st.session_state.count_reset}",
            help="Número del medidor eléctrico"
        )
        
        st.markdown("---")
        st.markdown("#### 📊 Datos de lectura y consumo")
        
        col_c, col_d, col_e = st.columns(3)
        
        with col_c:
            fecha_lectura_raw = st.text_input(
                "Fecha de acceso al medidor:",
                placeholder="08082025 o 08/08/2025",
                key=f"fecha_lectura_{st.session_state.count_reset}",
                help="Fecha en que se logró acceder al medidor"
            )
            fecha_lectura = formatear_fecha(fecha_lectura_raw)
            if fecha_lectura_raw and fecha_lectura != fecha_lectura_raw:
                st.caption(f"📅 Formateado: {fecha_lectura}")
            
            lectura_kwh = st.text_input(
                "Lectura registrada (kWh):",
                placeholder="20041",
                key=f"lectura_kwh_{st.session_state.count_reset}",
                help="Lectura total registrada en el medidor"
            )
        
        with col_d:
            consumo_total = st.text_input(
                "Consumo total (kWh):",
                placeholder="756",
                key=f"consumo_total_{st.session_state.count_reset}",
                help="Consumo total del periodo"
            )
            
            consumo_provisorio = st.text_input(
                "Consumo provisorio (kWh):",
                placeholder="184",
                key=f"consumo_provisorio_{st.session_state.count_reset}",
                help="Consumos provisorios descontados"
            )
        
        with col_e:
            fecha_inicio_periodo_raw = st.text_input(
                "Fecha inicio periodo apertura:",
                placeholder="11032025 o 11/03/2025",
                key=f"fecha_inicio_periodo_{st.session_state.count_reset}"
            )
            fecha_inicio_periodo = formatear_fecha(fecha_inicio_periodo_raw)
            if fecha_inicio_periodo_raw and fecha_inicio_periodo != fecha_inicio_periodo_raw:
                st.caption(f"📅 Formateado: {fecha_inicio_periodo}")
            
            fecha_fin_periodo_raw = st.text_input(
                "Fecha fin periodo apertura:",
                placeholder="08082025 o 08/08/2025",
                key=f"fecha_fin_periodo_{st.session_state.count_reset}"
            )
            fecha_fin_periodo = formatear_fecha(fecha_fin_periodo_raw)
            if fecha_fin_periodo_raw and fecha_fin_periodo != fecha_fin_periodo_raw:
                st.caption(f"📅 Formateado: {fecha_fin_periodo}")
        
        st.markdown("---")
        st.markdown("#### 💰 Reversa de Electricidad")
        
        monto_ajuste_input = st.text_input(
            "Monto reversa:",
            placeholder="39112",
            key=f"monto_ajuste_{st.session_state.count_reset}",
            help="Monto de la reversa de electricidad"
        )
        monto_ajuste = formatear_monto(monto_ajuste_input) if monto_ajuste_input else ""
        if monto_ajuste:
            st.caption(f"💰 Formateado: {monto_ajuste}")
        
        st.markdown("---")
        st.markdown("#### 📅 Historial de consumos (BO)")
        
        meses_historial = st.text_input(
            "Meses de historial:",
            placeholder="24",
            value="24",
            key=f"meses_historial_{st.session_state.count_reset}",
            help="Cantidad de meses del historial (ejemplo: 24)"
        )

elif tipo_carta == "aumento_consumo_halu_sinvisita":
    with st.expander("📊 DATOS ESPECÍFICOS - AUMENTO CONSUMO SIN VISITA", expanded=True):
        st.markdown("#### 📅 Historial de consumos")
        
        meses_historial_aumento = st.text_input(
            "Meses de historial:",
            placeholder="24",
            value="24",
            key=f"meses_historial_aumento_{st.session_state.count_reset}",
            help="Cantidad de meses del historial de consumo"
        )
        
        st.markdown("---")
        st.markdown("#### 💰 Rebaja aplicada")
        
        monto_rebaja_input = st.text_input(
            "Monto de rebaja:",
            placeholder="80058",
            key=f"monto_rebaja_{st.session_state.count_reset}",
            help="Monto de la rebaja aplicada por promedio histórico"
        )
        monto_rebaja = formatear_monto(monto_rebaja_input) if monto_rebaja_input else ""
        if monto_rebaja:
            st.caption(f"💰 Formateado: {monto_rebaja}")

elif tipo_carta == "aumento_consumo_nolu_sinvisita":
    with st.expander("📊 DATOS ESPECÍFICOS - AUMENTO CONSUMO NOLU SIN VISITA", expanded=True):
        st.markdown("#### 📅 Historial de consumos")
        
        meses_historial_nolu = st.text_input(
            "Meses de historial:",
            placeholder="24",
            value="24",
            key=f"meses_historial_nolu_{st.session_state.count_reset}",
            help="Cantidad de meses del historial de consumo"
        )

elif tipo_carta == "facturaciones_normalizadas":
    with st.expander("📊 DATOS ESPECÍFICOS - FACTURACIONES NORMALIZADAS", expanded=True):
        st.markdown("#### 📝 Motivo del reclamo")
        
        motivo_reclamo_fn = st.text_input(
            "Motivo del reclamo:",
            placeholder="error en la lectura",
            key=f"motivo_reclamo_fn_{st.session_state.count_reset}",
            help="Ejemplo: error en la lectura, aumento consumo, servicio no facturado, etc."
        )
        
        st.markdown("---")
        st.markdown("#### 📅 Rango de días de lectura")
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            dia_inicio_fn = st.text_input(
                "Día inicio:",
                placeholder="15",
                key=f"dia_inicio_fn_{st.session_state.count_reset}",
                help="Día de inicio del rango de lectura"
            )
        
        with col_b:
            dia_fin_fn = st.text_input(
                "Día fin:",
                placeholder="20",
                key=f"dia_fin_fn_{st.session_state.count_reset}",
                help="Día de fin del rango de lectura"
            )

elif tipo_carta == "carta_aporte_lectura":
    with st.expander("📊 DATOS ESPECÍFICOS - APORTE LECTURA", expanded=True):
        st.markdown("#### 📅 Fecha del requerimiento")
        
        fecha_requerimiento_raw = st.text_input(
            "Fecha del requerimiento:",
            placeholder="24112025 o 24/11/2025",
            key=f"fecha_requerimiento_{st.session_state.count_reset}",
            help="Fecha en que se efectuó el requerimiento"
        )
        fecha_requerimiento = formatear_fecha(fecha_requerimiento_raw)
        if fecha_requerimiento_raw and fecha_requerimiento != fecha_requerimiento_raw:
            st.caption(f"📅 Formateado: {fecha_requerimiento}")
        
        st.info("ℹ️ El N° de requerimiento se copiará automáticamente del N° GR")

elif tipo_carta == "error_lectura_regularizado_sgte_lectura":
    with st.expander("📊 DATOS ESPECÍFICOS - ERROR LECTURA REGULARIZADO", expanded=True):
        st.markdown("#### 📅 Rango de días de lectura")
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            dia_inicio_reg = st.text_input(
                "Día inicio:",
                placeholder="13",
                key=f"dia_inicio_reg_{st.session_state.count_reset}",
                help="Día de inicio del rango"
            )
        
        with col_b:
            dia_fin_reg = st.text_input(
                "Día fin:",
                placeholder="18",
                key=f"dia_fin_reg_{st.session_state.count_reset}",
                help="Día de fin del rango"
            )
        
        st.markdown("---")
        st.markdown("#### 📅 Historial de consumos (BO)")
        
        meses_historial_reg = st.text_input(
            "Meses de historial:",
            placeholder="24",
            value="24",
            key=f"meses_historial_reg_{st.session_state.count_reset}",
            help="Cantidad de meses del historial"
        )

elif tipo_carta == "error_lectura_nolu":
    with st.expander("📊 DATOS ESPECÍFICOS - ERROR LECTURA NOLU", expanded=True):
        st.markdown("#### 📅 Rango de días de lectura")
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            dia_inicio_nolu = st.text_input(
                "Día inicio:",
                placeholder="10",
                key=f"dia_inicio_nolu_{st.session_state.count_reset}",
                help="Día de inicio del rango"
            )
        
        with col_b:
            dia_fin_nolu = st.text_input(
                "Día fin:",
                placeholder="15",
                key=f"dia_fin_nolu_{st.session_state.count_reset}",
                help="Día de fin del rango"
            )
        
        st.markdown("---")
        st.markdown("#### 📊 Datos de facturación")
        
        col_c, col_d = st.columns(2)
        
        with col_c:
            fecha_factura_nolu_raw = st.text_input(
                "Fecha factura:",
                placeholder="15122025 o 15.12.2025",
                key=f"fecha_factura_nolu_{st.session_state.count_reset}",
                help="Fecha de la factura (se formateará con puntos: dd.mm.yyyy)"
            )
            # Para esta carta específica, formatear con PUNTOS en lugar de barras
            fecha_temp = formatear_fecha(fecha_factura_nolu_raw)
            fecha_factura_nolu = fecha_temp.replace('/', '.') if fecha_temp else ""
            if fecha_factura_nolu_raw and fecha_factura_nolu != fecha_factura_nolu_raw:
                st.caption(f"📅 Formateado: {fecha_factura_nolu}")
        
        with col_d:
            monto_factura_nolu_input = st.text_input(
                "Monto factura:",
                placeholder="36745",
                key=f"monto_factura_nolu_{st.session_state.count_reset}",
                help="Monto de la factura"
            )
            monto_factura_nolu = formatear_monto(monto_factura_nolu_input) if monto_factura_nolu_input else ""
            if monto_factura_nolu:
                st.caption(f"💰 Formateado: {monto_factura_nolu}")

elif tipo_carta == "atencion_emergencia_halu":
    with st.expander("📊 DATOS ESPECÍFICOS - ATENCIÓN EMERGENCIAS", expanded=True):
        st.markdown("#### 📅 Fechas de atención")
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            fecha_solicitud_emerg_raw = st.text_input(
                "Fecha solicitud:",
                placeholder="03102025 o 03/10/2025",
                key=f"fecha_solicitud_emerg_{st.session_state.count_reset}",
                help="Fecha de la solicitud de emergencia"
            )
            fecha_solicitud_emerg = formatear_fecha(fecha_solicitud_emerg_raw)
            if fecha_solicitud_emerg_raw and fecha_solicitud_emerg != fecha_solicitud_emerg_raw:
                st.caption(f"📅 Formateado: {fecha_solicitud_emerg}")
        
        with col_b:
            fecha_atencion_emerg_raw = st.text_input(
                "Fecha atención:",
                placeholder="15102025 o 15/10/2025",
                key=f"fecha_atencion_emerg_{st.session_state.count_reset}",
                help="Fecha en que se atendió la emergencia"
            )
            fecha_atencion_emerg = formatear_fecha(fecha_atencion_emerg_raw)
            if fecha_atencion_emerg_raw and fecha_atencion_emerg != fecha_atencion_emerg_raw:
                st.caption(f"📅 Formateado: {fecha_atencion_emerg}")
        
        st.markdown("---")
        st.markdown("#### 👤 Persona que solicitó atención emergencias")
        
        # ⭐ NUEVO: Selector de Señor/Señora para quien solicitó
        tratamiento_solicitante = st.selectbox(
            "Formalidad (Señor o Señora):",
            ["", "Señor", "Señora"],
            key=f"tratamiento_solicitante_{st.session_state.count_reset}",
            help="Tratamiento de quien solicitó la emergencia (puede ser distinto a quien va dirigida la carta)"
        )
        
        nombre_solicitante_raw = st.text_input(
            "Nombre completo:",
            placeholder="Mariana Lidia Espinoza Osorio",
            key=f"nombre_solicitante_{st.session_state.count_reset}",
            help="Nombre completo de quien solicitó la atención"
        )
        nombre_solicitante = capitalizar_texto(nombre_solicitante_raw)
        
        st.markdown("---")
        st.markdown("#### 📋 Orden de trabajo y nota de crédito")
        
        col_c, col_d = st.columns(2)
        
        with col_c:
            orden_trabajo = st.text_input(
                "N° Orden de trabajo:",
                placeholder="460506214",
                key=f"orden_trabajo_{st.session_state.count_reset}",
                help="Número de la orden de trabajo"
            )
        
        with col_d:
            nota_credito = st.text_input(
                "N° Nota de crédito:",
                placeholder="6669093",
                key=f"nota_credito_{st.session_state.count_reset}",
                help="Número de la nota de crédito"
            )
        
        st.markdown("---")
        st.markdown("#### 💰 Monto")
        
        monto_emerg_input = st.text_input(
            "Monto:",
            placeholder="24903",
            key=f"monto_emerg_{st.session_state.count_reset}",
            help="Monto de la atención de emergencia"
        )
        monto_emerg = formatear_monto(monto_emerg_input) if monto_emerg_input else ""
        if monto_emerg:
            st.caption(f"💰 Formateado: {monto_emerg}")

elif tipo_carta == "carta_falta_info":
    with st.expander("📊 DATOS ESPECÍFICOS - CARTA FALTA INFO", expanded=True):
        st.info("ℹ️ El N° de reclamo ingresado se copiará automáticamente del N° GR")
        st.markdown("Esta carta no requiere campos adicionales")

elif tipo_carta in ["carta_compromiso", "carta_compromiso_i5", "normal_avance"]:
    # Estas cartas no requieren campos específicos adicionales
    pass


# ================= TABLA EXCEL =================
with st.expander("📎 TABLA DE DATOS (Opcional - se agrega al final)"):
    st.warning("⚙️ **Funcionalidad en desarrollo** - Próximamente podrás adjuntar tablas dinámicas de Excel")
    st.info("💡 Copia y pega tabla desde Excel usando TAB como separador")
    tabla_excel = st.text_area(
        "Pegar datos aquí:",
        height=150,
        placeholder="Período\tConsumo\tMonto\nEne-2025\t150\t$45.000",
        key=f"tabla_excel_{st.session_state.count_reset}",
        disabled=True  # Deshabilitado mientras está en desarrollo
    )
    
    # TEMPORALMENTE DESHABILITADO - En desarrollo
    df = None
    # if tabla_excel:
    #     try:
    #         df = pd.read_csv(StringIO(tabla_excel), sep="\t")
    #         st.success(f"✅ {len(df)} filas cargadas")
    #         st.dataframe(df, use_container_width=True)
    #     except:
    #         st.error("⚠️ Formato incorrecto. Asegúrate de separar con TAB")

st.markdown("---")

# ================= BOTONES GENERAR Y DESCARGAR =================
col_btn1, col_btn2, col_btn3, col_btn4 = st.columns([1, 1.2, 1.2, 1])

with col_btn2:
    generar = st.button(
        "📝 GENERAR CARTA", 
        type="primary", 
        use_container_width=True,
        key=f"btn_generar_{st.session_state.count_reset}"
    )

with col_btn3:
    if st.session_state.get('carta_generada') and st.session_state.get('output_path'):
        if os.path.exists(st.session_state.output_path):
            with open(st.session_state.output_path, "rb") as f:
                st.download_button(
                    label="📥 DESCARGAR CARTA",
                    data=f,
                    file_name=os.path.basename(st.session_state.output_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"btn_descarga_{st.session_state.count_reset}"
                )

if generar:
    campos_requeridos = {
        "Nombre cliente": nombre_cliente,
        "N° GR": gr_numero,
        "Dirección": direccion,
        "Comuna": comuna,
        "Número cliente": numero_cliente,
        "Tratamiento": tratamiento
    }
    
    faltantes = [k for k, v in campos_requeridos.items() if not v]
    
    if faltantes:
        st.error(f"❌ Faltan campos obligatorios: {', '.join(faltantes)}")
    else:
        template_path = f"templates/{tipo_carta}.docx"
        
        if not os.path.exists(template_path):
            st.error(f"❌ No se encontró el template: {template_path}")
            st.info(f"📝 Copia tu carta Word a la carpeta 'templates/' como '{tipo_carta}.docx'")
        else:
            with st.spinner("⚡ Generando carta..."):
                try:
                    from datetime import datetime, timedelta, timezone
                    tz_chile = timezone(timedelta(hours=-3))
                    hoy = datetime.now(tz_chile)
                    
                    meses = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
                             "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
                    fecha_completa = f"{comuna}, {hoy.day} de {meses[hoy.month]} de {hoy.year}"
                    
                    if canal_ingreso == "WhatsApp":
                        texto_canal = "WhatsApp"
                    elif canal_ingreso in CANALES_EXTERNOS:
                        # Canales externos: solo el nombre (el template ya tiene "a través de")
                        texto_canal = canal_ingreso
                    elif canal_ingreso in CANALES_MASCULINO:
                        texto_canal = f"nuestro {canal_ingreso}"
                    else:
                        texto_canal = f"nuestra {canal_ingreso}"
                    
                    estimado = "Estimado" if tratamiento == "Señor" else "Estimada"
                    primer_nombre = nombre_cliente.split()[0] if nombre_cliente else "Cliente"
                    
                    doc = Document(template_path)
                    
                    # REEMPLAZOS COMUNES
                    reemplazos = {
                        "Valparaiso, 16 de diciembre de [202X]": fecha_completa,
                        "Valparaíso, 16 de diciembre de [202X]": fecha_completa,
                        "Valparaiso": comuna,
                        "Valparaíso": comuna,
                        "[Comuna]": comuna,
                        "DGR N.º XXXXXXX /[202X]": f"DGR N° {gr_numero} /{hoy.year}",
                        # Variantes de Ref.: Reclamo (se reemplazará después si hay caso SEC)
                        "Ref.: Reclamo N° 15965848": f"Ref.: Reclamo N° {gr_numero}",
                        "Ref.: Reclamo N° XXXXXXX": f"Ref.: Reclamo N° {gr_numero}",
                        "Reclamo N° XXXXXXX": f"Ref.: Reclamo N° {gr_numero}",  # Agrega "Ref.:" si falta
                        "Número de cliente: 15965848": f"Número de cliente: {numero_cliente}",
                        "[Señor(a)]": tratamiento,
                        "[Nombre y apellido reclamante]": nombre_cliente,
                        "[Dirección]": direccion,
                        "[Estimado(a) Nombre,]": f"{estimado} {primer_nombre},",
                        "[(Ej: nuestra Oficina Comercial / WhatsApp / App CGE 1Click / Call Center / Correo Electrónico / Página Web).]": f"{texto_canal}.",
                        "[Nombre y apellido Gerente Comercial]": GERENTES[zona],
                    }
                    
                    # Caso SEC/SERNAC (opcional)
                    # Si hay caso SEC/SERNAC, reemplaza COMPLETAMENTE la línea de Reclamo
                    # Si NO hay, deja solo "Ref.: Reclamo N°"
                    if tiene_caso_sec and caso_sec_numero and tipo_caso:
                        # Reemplazar COMPLETAMENTE con la línea de Caso SEC/SERNAC
                        texto_caso = f"Ref.: Caso {tipo_caso} N° {caso_sec_numero}, Reclamo N° {gr_numero}"
                        reemplazos["Ref.: Reclamo N° XXXXXXX"] = texto_caso
                        reemplazos["Reclamo N° XXXXXXX"] = texto_caso
                    # Si no hay caso SEC, ya se reemplazó arriba con "Ref.: Reclamo N° {gr_numero}"


                    # REEMPLAZOS ESPECÍFICOS POR TIPO DE CARTA
                    if tipo_carta == "error_lectura_halu":
                        if 'rango_lectura' in locals() and rango_lectura:
                            reemplazos["[XXXXXX y XXXXXX.]"] = f"{rango_lectura}."
                            reemplazos["[XXXXXX y XXXXXX]"] = rango_lectura
                            reemplazos["XXXXXX y XXXXXX"] = rango_lectura
                        if 'fecha_boleta' in locals() and fecha_boleta:
                            reemplazos["[día/mes/año]"] = fecha_boleta
                        if 'tipo_doc' in locals() and tipo_doc:
                            reemplazos["[boleta/factura]"] = tipo_doc
                        if 'numero_boleta' in locals() and numero_boleta:
                            reemplazos["[XXXXXX]"] = numero_boleta
                        if 'consumo_kwh' in locals() and consumo_kwh:
                            reemplazos["XXX kWh"] = f"{consumo_kwh} kWh"
                        if 'monto_boleta' in locals() and monto_boleta:
                            reemplazos["[$ XX.XXX]"] = monto_boleta
                    
                    elif tipo_carta == "apertura_casa_halu":
                        # Periodo sin acceso (con año inicio y fin)
                        if 'periodo_inicio' in locals() and periodo_inicio and 'periodo_fin' in locals() and periodo_fin and 'anio_inicio' in locals() and anio_inicio and 'anio_fin' in locals() and anio_fin:
                            # Si es el mismo año: "marzo a agosto del 2025"
                            if anio_inicio == anio_fin:
                                periodo_texto = f"{periodo_inicio} a {periodo_fin} del {anio_fin}"
                            else:
                                # Si son años diferentes: "marzo del 2024 a agosto del 2025"
                                periodo_texto = f"{periodo_inicio} del {anio_inicio} a {periodo_fin} del {anio_fin}"
                            reemplazos["[marzo a agosto del 2025]"] = periodo_texto
                        
                        # Número de medidor
                        if 'numero_medidor' in locals() and numero_medidor:
                            reemplazos["[E044124]"] = numero_medidor
                        
                        # Fecha de lectura
                        if 'fecha_lectura' in locals() and fecha_lectura:
                            reemplazos["[08/08/2025]"] = fecha_lectura
                        
                        # Lectura registrada
                        if 'lectura_kwh' in locals() and lectura_kwh:
                            reemplazos["[20.041]"] = lectura_kwh
                            reemplazos["[20041]"] = lectura_kwh
                        
                        # Consumo total
                        if 'consumo_total' in locals() and consumo_total:
                            reemplazos["[756]"] = consumo_total
                        
                        # Periodo completo
                        if 'fecha_inicio_periodo' in locals() and fecha_inicio_periodo and 'fecha_fin_periodo' in locals() and fecha_fin_periodo:
                            periodo_completo = f"{fecha_inicio_periodo} a {fecha_fin_periodo}"
                            reemplazos["[11/03/2025 a 08/08/2025]"] = periodo_completo
                        
                        # Consumo provisorio
                        if 'consumo_provisorio' in locals() and consumo_provisorio:
                            reemplazos["[184]"] = consumo_provisorio
                        
                        # Monto ajuste
                        if 'monto_ajuste' in locals() and monto_ajuste:
                            reemplazos["[$39.112]"] = monto_ajuste
                        
                        # Meses de historial
                        if 'meses_historial' in locals() and meses_historial:
                            reemplazos["[24]"] = meses_historial
                    
                    elif tipo_carta == "aumento_consumo_halu_sinvisita":
                        # Meses de historial
                        if 'meses_historial_aumento' in locals() and meses_historial_aumento:
                            reemplazos["[24]"] = meses_historial_aumento
                        
                        # Monto de rebaja
                        if 'monto_rebaja' in locals() and monto_rebaja:
                            reemplazos["[$80.058]"] = monto_rebaja
                    
                    elif tipo_carta == "aumento_consumo_nolu_sinvisita":
                        # Meses de historial
                        if 'meses_historial_nolu' in locals() and meses_historial_nolu:
                            reemplazos["[24]"] = meses_historial_nolu
                    
                    elif tipo_carta == "facturaciones_normalizadas":
                        # Motivo del reclamo
                        if 'motivo_reclamo_fn' in locals() and motivo_reclamo_fn:
                            reemplazos["[error en la lectura]"] = motivo_reclamo_fn
                        
                        # Rango de días
                        if 'dia_inicio_fn' in locals() and dia_inicio_fn and 'dia_fin_fn' in locals() and dia_fin_fn:
                            rango_dias = f"{dia_inicio_fn} y {dia_fin_fn}"
                            reemplazos["[15 y 20]"] = rango_dias
                    
                    elif tipo_carta == "carta_aporte_lectura":
                        # N° de requerimiento (igual que GR)
                        reemplazos["[15939748]"] = gr_numero
                        
                        # Fecha del requerimiento
                        if 'fecha_requerimiento' in locals() and fecha_requerimiento:
                            reemplazos["[24/11/2025]"] = fecha_requerimiento
                    
                    elif tipo_carta == "error_lectura_regularizado_sgte_lectura":
                        # Rango de días
                        if 'dia_inicio_reg' in locals() and dia_inicio_reg and 'dia_fin_reg' in locals() and dia_fin_reg:
                            rango_dias_reg = f"{dia_inicio_reg} y {dia_fin_reg}"
                            reemplazos["[13 y 18]"] = rango_dias_reg
                        
                        # Meses de historial
                        if 'meses_historial_reg' in locals() and meses_historial_reg:
                            reemplazos["[24]"] = meses_historial_reg
                    
                    elif tipo_carta == "error_lectura_nolu":
                        # Rango de días
                        if 'dia_inicio_nolu' in locals() and dia_inicio_nolu and 'dia_fin_nolu' in locals() and dia_fin_nolu:
                            rango_dias_nolu = f"{dia_inicio_nolu} y {dia_fin_nolu}"
                            reemplazos["[10 y 15]"] = rango_dias_nolu
                        
                        # Fecha factura
                        if 'fecha_factura_nolu' in locals() and fecha_factura_nolu:
                            reemplazos["[15.12.2025]"] = fecha_factura_nolu
                        
                        # Lectura kWh (SIN signo $, solo número con puntos)
                        if 'monto_factura_nolu_input' in locals() and monto_factura_nolu_input:
                            # Formatear solo con puntos de miles, sin $
                            valor_limpio = str(monto_factura_nolu_input).replace('$', '').replace('.', '').replace(',', '').strip()
                            if valor_limpio:
                                numero = int(valor_limpio)
                                lectura_formateada = f"{numero:,}".replace(',', '.')
                                reemplazos["[$65.000]"] = lectura_formateada
                                reemplazos["[36.745]"] = lectura_formateada
                    
                    elif tipo_carta == "atencion_emergencia_halu":
                        # Fechas
                        if 'fecha_solicitud_emerg' in locals() and fecha_solicitud_emerg:
                            reemplazos["[03/10/2025]"] = fecha_solicitud_emerg
                        
                        if 'fecha_atencion_emerg' in locals() and fecha_atencion_emerg:
                            reemplazos["[15/10/2025]"] = fecha_atencion_emerg
                        
                        # Persona solicitante con "el Sr." o "la Sra."
                        if 'nombre_solicitante' in locals() and nombre_solicitante and 'tratamiento_solicitante' in locals() and tratamiento_solicitante:
                            # Determinar artículo según tratamiento DEL SOLICITANTE (no del destinatario)
                            if tratamiento_solicitante == "Señor":
                                texto_solicitante = f"el Sr. {nombre_solicitante}"
                            else:  # Señora
                                texto_solicitante = f"la Sra. {nombre_solicitante}"
                            
                            reemplazos["[Mariana Lidia Espinoza Osorio]"] = nombre_solicitante
                            reemplazos["[el(a) Sr(a). XXXXXX XXXXXX]"] = texto_solicitante
                        
                        # Nota de crédito
                        if 'nota_credito' in locals() and nota_credito:
                            reemplazos["[6669093]"] = nota_credito
                        
                        # Orden de trabajo
                        if 'orden_trabajo' in locals() and orden_trabajo:
                            reemplazos["[460506214]"] = orden_trabajo
                        
                        # Monto
                        if 'monto_emerg' in locals() and monto_emerg:
                            reemplazos["[$24.903]"] = monto_emerg
                    
                    elif tipo_carta == "carta_falta_info":
                        # N° de reclamo ingresado (igual que GR)
                        reemplazos["[15939815]"] = gr_numero
                    
                    reemplazos_ordenados = sorted(reemplazos.items(), key=lambda x: len(x[0]), reverse=True)
                    
                    def aplicar_reemplazos(texto):
                        for key, value in reemplazos_ordenados:
                            texto = texto.replace(key, str(value))
                        return texto
                    
                    # Aplicar a párrafos
                    for paragraph in doc.paragraphs:
                        texto = paragraph.text
                        texto_nuevo = aplicar_reemplazos(texto)
                        if texto_nuevo != texto:
                            paragraph.clear()
                            run = paragraph.add_run(texto_nuevo)
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            
                            # SOLO estas palabras van en negrita (SIN fecha, SIN DGR)
                            palabras_con_negrita = [
                                tratamiento, 
                                nombre_cliente, 
                                direccion,
                                "Número de cliente:", 
                                "Ref.: Reclamo N°", 
                                GERENTES[zona], 
                                "COMPAÑÍA GENERAL DE ELECTRICIDAD S.A."
                            ]
                            
                            # Aplica negrita SOLO si coincide Y NO es fecha ni DGR
                            es_fecha = (f"{comuna}," in texto_nuevo and "de" in texto_nuevo and str(hoy.year) in texto_nuevo)
                            es_dgr = ("DGR N°" in texto_nuevo and str(hoy.year) in texto_nuevo)
                            
                            if any(x in texto_nuevo for x in palabras_con_negrita) and not es_fecha and not es_dgr:
                                run.font.bold = True
                            else:
                                run.font.bold = False
                            
                            # Excepción: Si el párrafo es SOLO la comuna (una línea), SÍ va en negrita
                            if texto_nuevo.strip() == comuna:
                                run.font.bold = True
                    
                    # Tablas
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    texto = aplicar_reemplazos(p.text)
                                    if texto != p.text:
                                        p.clear()
                                        r = p.add_run(texto)
                                        r.font.name = 'Arial'
                                        r.font.size = Pt(10)
                    
                    # Excel Anexo
                    if df is not None:
                        doc.add_page_break()
                        doc.add_heading("Anexo - Datos Adicionales", level=1)
                        table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                        table.style = "Light Grid Accent 1"
                        for i, col in enumerate(df.columns):
                            cell = table.rows[0].cells[i]
                            cell.text = str(col)
                        for i, row in df.iterrows():
                            for j, value in enumerate(row):
                                table.rows[i+1].cells[j].text = str(value)
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"Carta_{tipo_carta}_{gr_numero}_{timestamp}.docx"
                    output_path = os.path.join("output", filename)
                    doc.save(output_path)
                    
                    st.session_state.carta_generada = True
                    st.session_state.output_path = output_path
                    
                    st.session_state.vista_previa_html = f"""
                        <div style='font-family:Arial;font-size:10pt;background:white;padding:20px;border:1px solid #ddd;border-radius:8px;color:black'>
                            <div style='text-align:right;margin-bottom:20px'>
                                <p style='margin:0'>{comuna}, {hoy.day} de {meses[hoy.month]} de {hoy.year}</p>
                                <p style='margin:0'>DGR N° {gr_numero} /{hoy.year}</p>
                            </div>
                            <p><strong>{tratamiento}</strong></p>
                            <p><strong>{nombre_cliente}</strong></p>
                            <p><strong>{direccion}</strong></p>
                            <p><strong>{comuna}</strong></p>
                            <p style='margin-top:10px'><strong>Número de cliente: {numero_cliente}</strong></p>
                            <p><strong>Ref.: Reclamo N° {gr_numero}</strong></p>
                            <p style='margin-top:15px'>{estimado} {primer_nombre},</p>
                            <p style='text-align:justify'>Junto con saludar, le confirmamos que hemos recibido su reclamo...</p>
                        </div>
                    """
                    st.success("✅ ¡Carta generada exitosamente!")
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
                    with st.expander("🔍 Ver detalles del error"):
                        st.code(str(e))

# MOSTRAR VISTA PREVIA
if st.session_state.get('carta_generada') and 'vista_previa_html' in st.session_state:
    with st.expander("👁️ VER VISTA PREVIA", expanded=True):
        st.markdown(st.session_state.vista_previa_html, unsafe_allow_html=True)

# FOOTER
st.markdown("---")
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("📁 Cartas", len([f for f in os.listdir('output') if f.endswith('.docx')]))
with col2:
    st.metric("📄 Templates", len([f for f in os.listdir('templates') if f.endswith('.docx')]))
with col3:
    st.metric("⏱️ Generación", "< 60 seg")

st.markdown(f"""
<div style='text-align:center;padding:20px;color:#64748b;font-size:0.9em'>
    <p style='margin:5px 0'>⚡ <strong>Automatizador de Cartas</strong></p>
    <p style='margin:5px 0'>Desarrollado por <a href='https://ciberbyte.vercel.app/' target='_blank' style='color:#0ea5e9; text-decoration:none; font-weight:bold;'>CiberByte</a> <span style='color:#94a3b8'>/</span> <a href='https://wa.me/56979693753?text=Hola%20Javier,%20consulta%20sobre%20el%20Automatizador%20de%20Cartas' target='_blank' style='color:#1e40af; text-decoration:none; font-weight:600;'>Javier Ruiz Arismendi <svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" viewBox="0 0 24 24" fill="#25D366" style="vertical-align: middle; margin-left: 3px;"><path d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347m-5.421 7.403h-.004a9.87 9.87 0 01-5.031-1.378l-.361-.214-3.741.982.998-3.648-.235-.374a9.86 9.86 0 01-1.51-5.26c.001-5.45 4.436-9.884 9.888-9.884 2.64 0 5.122 1.03 6.988 2.898a9.825 9.825 0 012.893 6.994c-.003 5.45-4.437 9.884-9.885 9.884m8.413-18.297A11.815 11.815 0 0012.05 0C5.495 0 .16 5.335.157 11.892c0 2.096.547 4.142 1.588 5.945L.057 24l6.305-1.654a11.882 11.882 0 005.683 1.448h.005c6.554 0 11.89-5.335 11.893-11.893a11.821 11.821 0 00-3.48-8.413Z"/></svg></a></p>
    <p style='margin:5px 0'>© 2025 - Sistema de Generación Automática</p>
</div>
""", unsafe_allow_html=True)